from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from docx import Document
import pandas as pd
import os
import re
from typing import Dict, List, Tuple
from datetime import datetime
from enum import Enum

class DatabaseDialect(Enum):
    MYSQL = "mysql"
    POSTGRESQL = "postgresql"
    SQLITE = "sqlite"
    MSSQL = "mssql"

class SQLTypeMapper:
    def __init__(self, dialect: DatabaseDialect):
        self.dialect = dialect
        
    def get_integer_type(self, max_val: int) -> str:
        if self.dialect == DatabaseDialect.POSTGRESQL:
            if max_val < 32767:
                return 'SMALLINT'
            elif max_val < 2147483647:
                return 'INTEGER'
            return 'BIGINT'
        elif self.dialect == DatabaseDialect.MYSQL:
            if max_val < 32767:
                return 'SMALLINT'
            elif max_val < 2147483647:
                return 'INT'
            return 'BIGINT'
        return 'INTEGER'  # Default for SQLite and others
        
    def get_decimal_type(self, precision: int = 18, scale: int = 2) -> str:
        if self.dialect == DatabaseDialect.POSTGRESQL:
            return f'NUMERIC({precision},{scale})'
        return f'DECIMAL({precision},{scale})'
        
    def get_varchar_type(self, length: int) -> str:
        if self.dialect == DatabaseDialect.POSTGRESQL:
            if length > 255:
                return 'TEXT'
            return f'VARCHAR({length})'
        elif self.dialect == DatabaseDialect.MYSQL:
            if length > 65535:
                return 'LONGTEXT'
            elif length > 16383:
                return 'MEDIUMTEXT'
            return f'VARCHAR({length})'
        return f'VARCHAR({length})'  # Default for others
        
    def get_date_type(self) -> str:
        if self.dialect == DatabaseDialect.POSTGRESQL:
            return 'TIMESTAMP'
        elif self.dialect == DatabaseDialect.MYSQL:
            return 'DATETIME'
        return 'TIMESTAMP'

def analyze_column_data(df: pd.DataFrame, dialect: DatabaseDialect = DatabaseDialect.POSTGRESQL) -> Dict[str, Tuple[str, Dict]]:
    """
    Analyze DataFrame columns and suggest SQL data types with additional metadata.
    Returns a dictionary of column names and their suggested SQL types plus metadata.
    """
    sql_types = {}
    type_mapper = SQLTypeMapper(dialect)
    
    for column in df.columns:
        # Clean column name for SQL
        clean_col_name = re.sub(r'\W+', '_', str(column)).lower().strip('_')
        
        # Get non-null values
        non_null_values = df[column].dropna()
        null_count = df[column].isna().sum()
        total_count = len(df[column])
        
        metadata = {
            'original_name': column,
            'null_count': null_count,
            'total_count': total_count,
            'null_percentage': (null_count / total_count) * 100 if total_count > 0 else 0
        }
        
        if len(non_null_values) == 0:
            sql_types[clean_col_name] = (type_mapper.get_varchar_type(255), metadata)
            continue
            
        # Check if all values are numeric
        try:
            numeric_values = pd.to_numeric(non_null_values)
            metadata['min_value'] = float(numeric_values.min())
            metadata['max_value'] = float(numeric_values.max())
            
            if all(numeric_values.astype(int) == numeric_values):
                max_val = numeric_values.max()
                sql_type = type_mapper.get_integer_type(max_val)
            else:
                # Analyze decimal places
                decimal_places = max(
                    len(str(x).split('.')[-1]) if '.' in str(x) else 0 
                    for x in non_null_values
                )
                metadata['decimal_places'] = decimal_places
                sql_type = type_mapper.get_decimal_type(18, decimal_places)
            
            sql_types[clean_col_name] = (sql_type, metadata)
            continue
        except:
            pass
        
        # Check if all values are dates
        try:
            dates = pd.to_datetime(non_null_values)
            metadata['min_date'] = dates.min().strftime('%Y-%m-%d')
            metadata['max_date'] = dates.max().strftime('%Y-%m-%d')
            sql_types[clean_col_name] = (type_mapper.get_date_type(), metadata)
            continue
        except:
            pass
        
        # Text analysis
        str_lengths = non_null_values.astype(str).str.len()
        metadata['min_length'] = int(str_lengths.min())
        metadata['max_length'] = int(str_lengths.max())
        metadata['avg_length'] = float(str_lengths.mean())
        
        # Get appropriate VARCHAR length with 20% padding
        max_length = int(str_lengths.max() * 1.2)
        sql_types[clean_col_name] = (type_mapper.get_varchar_type(max_length), metadata)
    
    return sql_types

def generate_sql_schema(excel_path: str, table_name: str, dialect: DatabaseDialect = DatabaseDialect.POSTGRESQL) -> Tuple[str, Dict]:
    """
    Generate SQL schema from Excel file with detailed column analysis.
    Returns both the SQL schema and column metadata.
    """
    # Read Excel file
    df = pd.read_excel(excel_path)
    
    # Analyze columns
    column_analysis = analyze_column_data(df, dialect)
    
    # Generate CREATE TABLE statement
    create_table = f"CREATE TABLE {table_name} (\n"
    columns = []
    
    for col, (sql_type, metadata) in column_analysis.items():
        nullable = "NULL" if metadata['null_count'] > 0 else "NOT NULL"
        columns.append(f"    {col} {sql_type} {nullable}")
    
    create_table += ",\n".join(columns)
    create_table += "\n);"
    
    # Save the schema and analysis to files
    base_path = os.path.dirname(excel_path)
    
    # Save SQL schema
    schema_path = os.path.join(base_path, f'schema_{dialect.value}.sql')
    with open(schema_path, 'w') as f:
        f.write(create_table)
    
    # Save detailed analysis as JSON
    import json
    analysis_path = os.path.join(base_path, 'column_analysis.json')
    analysis_dict = {col: metadata for col, (_, metadata) in column_analysis.items()}
    with open(analysis_path, 'w') as f:
        json.dump(analysis_dict, f, indent=2)
    
    return create_table, column_analysis

def copy_text_and_style(worksheet, cell_address, paragraph):
    """
    Copy text and basic style from a Word paragraph to an Excel cell.
    """
    cell = worksheet[cell_address]
    cell.value = paragraph.text

    # Apply basic formatting
    cell.font = Font(name='Calibri', size=11)
    if paragraph.style.font.bold:
        cell.font = Font(bold=True)
    if paragraph.style.font.italic:
        cell.font = Font(italic=True)
    cell.alignment = Alignment(horizontal='left', wrap_text=True)

def copy_table_to_excel(worksheet, start_row, table):
    """
    Copy a Word table to an Excel worksheet starting at the specified row.
    """
    current_row = start_row
    for row in table.rows:
        current_col = 1
        for cell in row.cells:
            excel_cell = worksheet.cell(row=current_row, column=current_col)
            excel_cell.value = cell.text
            excel_cell.alignment = Alignment(horizontal='left', wrap_text=True)
            current_col += 1
        current_row += 1
    return current_row

def export_word_to_excel(word_path, excel_path):
    """
    Export content from a Word document to an Excel file.
    """
    # Load Word document
    doc = Document(word_path)

    # Create Excel workbook and worksheet
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "WordToExcel"

    row = 1

    # Loop through elements in the Word document
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            try:
                # Safely access paragraph if it exists
                paragraph = doc.paragraphs[row - 1] if row - 1 < len(doc.paragraphs) else None
                if paragraph and paragraph.text.strip():  # Ensure paragraph exists and is not empty
                    cell_address = f"A{row}"
                    copy_text_and_style(worksheet, cell_address, paragraph)
                    row += 1
            except IndexError as e:
                print(f"Skipping invalid paragraph at index {row - 1}: {e}")
        elif element.tag.endswith('tbl'):  # Table
            # Get the table object
            try:
                table = next((t for t in doc.tables if t._element == element), None)
                if table:
                    row = copy_table_to_excel(worksheet, row, table)
            except Exception as e:
                print(f"Skipping invalid table: {e}")

    # Ensure the output directory exists
    output_dir = os.path.dirname(excel_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save Excel file
    workbook.save(excel_path)

# Paths
word_file_path = "/Users/mutai/Desktop/KERICHO_COUNTY_FINANCE.docx"
excel_file_path = "output/WordToExcel.xlsx"

# Export the Word document to Excel
export_word_to_excel(word_file_path, excel_file_path)

# Generate SQL schema for PostgreSQL (default)
table_name = "kericho_county_finance"
sql_schema, column_analysis = generate_sql_schema(excel_file_path, table_name)

print("\nGenerated PostgreSQL Schema:")
print(sql_schema)

# Also generate MySQL version
mysql_schema, _ = generate_sql_schema(excel_file_path, table_name, DatabaseDialect.MYSQL)
print("\nGenerated MySQL Schema:")
print(mysql_schema)

print("\nDetailed column analysis has been saved to 'column_analysis.json'")
